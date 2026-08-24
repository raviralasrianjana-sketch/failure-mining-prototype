"""
pipeline.py
-----------
The "brain" of the project. This file has NO dashboard code in it --
it's pure data processing so it's easy to test and reason about on its own.

Stages (matches the flow in the problem statement):
  1. load_data      -> read the CSV
  2. sanitize_text   -> remove PII (no names/phone/email should reach the model)
  3. vectorize_text  -> turn each note into numbers (TF-IDF "embedding")
  4. cluster_notes    -> group similar notes together (KMeans)
  5. label_clusters   -> give each cluster a human-readable name (top keywords)
  6. build_trend_table -> count failures per model per month
  7. build_insights    -> plain-English summary of the top issues

WHY TF-IDF INSTEAD OF A NEURAL EMBEDDING MODEL?
Neural sentence embeddings (e.g. sentence-transformers, OpenAI embeddings)
would work too and are mentioned in the original brief. But they need a
model download / API key and a GPU helps. For a hackathon MVP, TF-IDF:
  - runs instantly, offline, no API key
  - is easy to explain in a demo ("it counts important words in each note
    and compares notes based on which words they share")
  - is a legitimate, commonly used baseline for text clustering

If you have time near the end of the hackathon, swapping in
sentence-transformers or Azure OpenAI embeddings is a small, isolated change
-- only vectorize_text() needs to change. Everything downstream (clustering,
labeling, trends) stays the same. That's a good "future work" slide!
"""

import re
import io
import os
from datetime import datetime

import numpy as np
import pandas as pd
from sklearn.feature_extraction.text import TfidfVectorizer
from sklearn.cluster import KMeans
from sklearn.metrics import silhouette_score
from sklearn.metrics.pairwise import cosine_similarity

from review_preprocessor import is_structured_failure_data, preprocess_raw_reviews


# ---------------------------------------------------------------------------
# 0. EXTRACT TEXT FROM UNSTRUCTURED FILES (PDF / Word / images)
# ---------------------------------------------------------------------------
# CSV/Excel/JSON already come as neat rows and columns. PDFs, Word docs, and
# images (e.g. a photographed service note or a screenshot) don't -- they're
# just "a page of text" or "a picture of text". So for these, we first pull
# out the raw text, then PARSE that text into the same table shape
# (product_model, date, symptom_text, fix_text) the rest of the pipeline
# expects. This parsing uses simple label-matching (e.g. a line starting
# with "Model:" or "Fix:") -- it works well for notes written with those
# labels, and falls back to sensible defaults ("Unknown" model, today's
# date) when a line doesn't have them, so nothing crashes on messy input.

def extract_text_from_pdf(path_or_buffer) -> str:
    """Pulls all text out of a PDF, page by page."""
    import pdfplumber
    text_parts = []
    with pdfplumber.open(path_or_buffer) as pdf:
        for page in pdf.pages:
            page_text = page.extract_text() or ""
            text_parts.append(page_text)
    return "\n".join(text_parts)


def extract_text_from_docx(path_or_buffer) -> str:
    """Pulls all paragraph text out of a Word document."""
    import docx
    document = docx.Document(path_or_buffer)
    return "\n".join(p.text for p in document.paragraphs)


def extract_text_from_image(path_or_buffer) -> str:
    """
    OCR (Optical Character Recognition): reads the pixels of an image and
    recognizes any text in it -- this is how we handle scanned notes or
    screenshots. Uses Tesseract OCR under the hood via pytesseract.

    IMPORTANT: pytesseract is just a Python wrapper -- it needs the actual
    Tesseract OCR *program* installed separately on the computer (it's not
    something pip can install by itself). If it's missing, we raise a
    clear error telling the user how to fix it, instead of a cryptic crash.
    """
    import pytesseract
    from PIL import Image
    image = Image.open(path_or_buffer)
    try:
        return pytesseract.image_to_string(image)
    except pytesseract.TesseractNotFoundError:
        raise ValueError(
            "Image text reading needs the Tesseract OCR program installed "
            "on this computer (separate from the Python packages). "
            "On Windows: download and install it from "
            "https://github.com/UB-Mannheim/tesseract/wiki, then restart "
            "the app. On Mac: run 'brew install tesseract'. On Linux: run "
            "'sudo apt install tesseract-ocr'."
        )


def extract_text_from_plain_file(path_or_buffer) -> str:
    """Read a text export without assuming a particular text encoding."""
    if hasattr(path_or_buffer, "seek"):
        path_or_buffer.seek(0)
    raw = path_or_buffer.read() if hasattr(path_or_buffer, "read") else open(path_or_buffer, "rb").read()
    if isinstance(raw, str):
        return raw
    for encoding in ("utf-8-sig", "utf-8", "cp1252", "latin-1"):
        try:
            return raw.decode(encoding)
        except UnicodeDecodeError:
            continue
    return raw.decode("utf-8", errors="replace")


# Regex patterns used to spot labeled fields inside freeform text, e.g.
# a line like "Model: B200   Date: 2025-11-02" or "Fix: replaced fan".
_MODEL_RE = re.compile(r"(?:model|product)\s*[:\-]\s*([A-Za-z0-9\-\_]+)", re.IGNORECASE)
_DATE_RE = re.compile(r"(\d{4}-\d{1,2}-\d{1,2}|\d{1,2}[/\-]\d{1,2}[/\-]\d{2,4})")
_FIX_RE = re.compile(r"(?:fix|resolution|repair(?:ed)?)\s*[:\-]\s*(.+)", re.IGNORECASE)


def parse_freeform_text_to_notes(raw_text: str, source_filename: str = "uploaded_file") -> pd.DataFrame:
    """
    Turns a blob of raw extracted text into a table of service notes.

    Strategy (kept simple on purpose -- this is a heuristic, not magic):
    - Split the text into non-empty lines/paragraphs.
    - A line matching "Fix:"/"Resolution:"/"Repair:" is treated as the FIX
      for the note directly above it, not a new note.
    - A line that is ONLY a "Model: X  Date: Y" header (no other text) is
      remembered and applied to the NEXT real note line, rather than
      becoming an empty note of its own -- this handles the common style
      where the model/date sits on its own line above the symptom text.
    - Every other non-empty line becomes its own note, using any
      model/date found on that same line (or the pending header, if one
      was just seen).
    - Anything not found falls back to "Unknown" model / today's date, so
      the pipeline never crashes on messy, real-world text.
    """
    lines = [ln.strip() for ln in raw_text.splitlines() if ln.strip()]

    notes = []
    pending_model = None
    pending_date = None

    for line in lines:
        fix_match = _FIX_RE.search(line)
        if fix_match and notes:
            # Attach this as the fix for the most recent note instead of
            # creating a brand new one.
            notes[-1]["fix_text"] = fix_match.group(1).strip()
            continue

        model_match = _MODEL_RE.search(line)
        date_match = _DATE_RE.search(line)

        # Strip out the "Model: ..." / "Date: ..." labels to see what's left.
        # (We remove the "date" label word separately from the date VALUE,
        # since _DATE_RE only matches the value itself, e.g. "2026-01-15",
        # not the word "Date:" in front of it.)
        remainder = _MODEL_RE.sub("", line)
        remainder = re.sub(r"date\s*[:\-]\s*", "", remainder, flags=re.IGNORECASE)
        remainder = _DATE_RE.sub("", remainder).strip(" :-,\t")

        if (model_match or date_match) and not remainder:
            # This line is ONLY a header (e.g. "Model: A100  Date: 2026-01-15")
            # -- remember it for the next real note line instead of creating
            # an empty note now.
            if model_match:
                pending_model = model_match.group(1)
            if date_match:
                pending_date = date_match.group(1)
            continue

        product_model = (
            model_match.group(1) if model_match
            else pending_model if pending_model
            else "Unknown"
        )
        raw_date = date_match.group(1) if date_match else pending_date
        if raw_date:
            try:
                note_date = pd.to_datetime(raw_date)
            except (ValueError, TypeError):
                note_date = pd.NaT
        else:
            note_date = pd.NaT

        symptom = remainder if remainder else line

        notes.append({
            "product_model": product_model,
            "date": note_date,
            "symptom_text": symptom,
            "fix_text": "",
        })

        # A pending header only applies to the next note that follows it.
        pending_model, pending_date = None, None

    if not notes:
        raise ValueError(
            f"Couldn't find any readable text in '{source_filename}'. "
            "If this is a scanned image, make sure the text is clear and not too small."
        )

    df = pd.DataFrame(notes)
    df.insert(0, "note_id", [f"{source_filename}-{i+1:04d}" for i in range(len(df))])
    return df


# ---------------------------------------------------------------------------
# 1. LOAD DATA
# ---------------------------------------------------------------------------
def load_data(path_or_buffer, filename: str = None) -> pd.DataFrame:
    """
    Loads service notes from CSV, Excel (.xlsx/.xls), TSV, JSON, PDF, Word
    (.docx), or images (.png/.jpg/.jpeg).

    Why support multiple formats?
    CSV is the simplest and most universal (plain text, works everywhere),
    which is why we started with it. But in practice, teams often keep
    their service logs in Excel, or even as scanned PDFs/photos of paper
    notes -- so we detect the file type from its name/extension and use
    the matching reader. Everything AFTER this function (sanitizing,
    clustering, etc.) doesn't care what the original file format was --
    it all becomes the same kind of table (a pandas DataFrame) at this
    point.

    Args:
        path_or_buffer: a file path (str) OR an in-memory uploaded file
            (e.g. from Streamlit's file_uploader), which behaves like a
            file object rather than a path.
        filename: the original filename, used to detect format when
            path_or_buffer doesn't carry a reliable extension on its own
            (e.g. a temp file already renamed to .csv).
    """
    name = os.path.basename(filename or getattr(path_or_buffer, "name", None) or str(path_or_buffer))
    ext = name.lower().rsplit(".", 1)[-1] if "." in name else ""

    # detection_info is only ever set for the CSV/Excel/TSV/JSON branch,
    # and only describes anything interesting when raw review data was
    # detected and converted. app.py reads this (via df.attrs) to show a
    # small "Detected input type: ..." caption. It's None for every input
    # that already matched the existing structured/PDF/DOCX/image paths.
    detection_info = None

    # Streamlit uploads and BytesIO objects are sometimes reused after a
    # reader has consumed them.  Always rewind before handing them to a
    # format-specific reader.
    if hasattr(path_or_buffer, "seek"):
        path_or_buffer.seek(0)

    if ext in ("xlsx", "xls", "xlsm", "ods"):
        df = pd.read_excel(path_or_buffer)
        df, detection_info = _route_structured_or_reviews(df, name)
    elif ext == "tsv":
        df = pd.read_csv(path_or_buffer, sep="\t")
        df, detection_info = _route_structured_or_reviews(df, name)
    elif ext == "json":
        df = pd.read_json(path_or_buffer)
        df, detection_info = _route_structured_or_reviews(df, name)
    elif ext == "pdf":
        raw_text = extract_text_from_pdf(path_or_buffer)
        df = parse_freeform_text_to_notes(raw_text, source_filename=name)
    elif ext == "docx":
        raw_text = extract_text_from_docx(path_or_buffer)
        df = parse_freeform_text_to_notes(raw_text, source_filename=name)
    elif ext in ("png", "jpg", "jpeg"):
        raw_text = extract_text_from_image(path_or_buffer)
        df = parse_freeform_text_to_notes(raw_text, source_filename=name)
    elif ext in ("webp", "gif", "bmp", "tif", "tiff"):
        raw_text = extract_text_from_image(path_or_buffer)
        df = parse_freeform_text_to_notes(raw_text, source_filename=name)
    elif ext in ("txt", "md", "log"):
        raw_text = extract_text_from_plain_file(path_or_buffer)
        df = parse_freeform_text_to_notes(raw_text, source_filename=name)
    else:
        # CSV has a few common filename variants.  Do not silently interpret
        # an arbitrary binary upload as CSV: that produces confusing parser
        # errors and can make a valid upload look broken.
        if ext not in ("csv", ""):
            raise ValueError(
                f"'{name}' is not a supported input format. "
                "Supported formats: CSV, TSV, Excel, JSON, PDF, Word, text, and images."
            )
        try:
            df = pd.read_csv(path_or_buffer)
        except UnicodeDecodeError:
            if hasattr(path_or_buffer, "seek"):
                path_or_buffer.seek(0)
            df = pd.read_csv(path_or_buffer, encoding="latin-1")
        df, detection_info = _route_structured_or_reviews(df, name)

    # Required columns check -- fail with a clear message rather than a
    # confusing crash deeper in the pipeline if the file is missing a column.
    # Dates are optional: clustering and drill-down still work for text-only
    # exports, while temporal views gracefully disable themselves.
    required = {"product_model", "symptom_text"}
    missing = required - set(df.columns)
    if missing:
        raise ValueError(
            f"Uploaded file is missing required column(s): {', '.join(sorted(missing))}. "
            f"Expected at least: {', '.join(sorted(required))}."
        )

    if "date" not in df.columns:
        df["date"] = pd.NaT
    else:
        df["date"] = pd.to_datetime(df["date"], errors="coerce")
    if "fix_text" not in df.columns:
        df["fix_text"] = ""

    # --- serial_range / serial_number (metadata field from the brief) ---
    # Accept either column name -- some exports call it "serial_range"
    # (a batch/range identifier), others "serial_number" (one unit's own
    # serial). We normalize both into a single "serial_range" column so
    # the rest of the app only has to deal with one name. If neither is
    # present, we don't fail the whole upload over an optional field --
    # we just fill in "Unknown" so downstream grouping still works.
    if "serial_range" in df.columns:
        pass
    elif "serial_number" in df.columns:
        df["serial_range"] = df["serial_number"]
    else:
        df["serial_range"] = "Unknown"
    df["serial_range"] = df["serial_range"].fillna("Unknown").astype(str)

    # --- optional structured tag (e.g. "no_power", "overheating") ---
    # This is the brief's "Optional structured tags" field. It's allowed
    # to be missing or blank on any given row (real data is inconsistent)
    # -- we just normalize it to an empty string rather than NaN so it's
    # easy to check "does this row have a tag?" downstream.
    if "tag" not in df.columns:
        df["tag"] = ""
    df["tag"] = df["tag"].fillna("").astype(str).str.strip()

    # Stash detection info as DataFrame metadata (survives .copy() and
    # column assignment, which is all the rest of the pipeline does to
    # this object) so app.py can show what kind of input was detected,
    # without needing load_data()'s return type/signature to change.
    df.attrs["detection_info"] = detection_info or {"input_type": "structured"}

    return df


def _route_structured_or_reviews(df: pd.DataFrame, filename: str):
    """
    Decides whether a CSV/Excel/TSV/JSON upload is already a structured
    Failure Mining dataset (goes straight through, UNCHANGED) or raw
    review data (goes through review_preprocessor.py first).

    This is the ONLY new branching point introduced for raw review
    support. If the required columns are already present, this function
    returns the DataFrame completely untouched -- guaranteeing identical
    behavior to before this feature existed, for any dataset that
    already worked.
    """
    if is_structured_failure_data(df):
        return df, {"input_type": "structured"}

    standardized_df, detection_info = preprocess_raw_reviews(df, source_filename=filename)
    return standardized_df, detection_info


# ---------------------------------------------------------------------------
# 2. SANITIZE (remove PII) -- a Constraint & Guardrail from the brief
# ---------------------------------------------------------------------------
# Simple regex patterns. Not perfect (real PII scrubbing is a deep topic),
# but demonstrates the *principle*: no email/phone/name should reach the
# model or the dashboard.
EMAIL_RE = re.compile(r"[\w\.\-]+@[\w\.\-]+\.\w+")
PHONE_RE = re.compile(r"(\+?\d[\d\-\s]{8,}\d)")
# Very simple "Customer <First Last>" name pattern used by our fake data.
NAME_RE = re.compile(r"Customer [A-Z][a-z]+ [A-Z][a-z]+")


def sanitize_text(text: str) -> str:
    """Redact emails, phone numbers, and simple name patterns from a note."""
    if not isinstance(text, str):
        return ""
    text = EMAIL_RE.sub("[REDACTED_EMAIL]", text)
    text = PHONE_RE.sub("[REDACTED_PHONE]", text)
    text = NAME_RE.sub("Customer [REDACTED_NAME]", text)
    return text


def sanitize_dataframe(df: pd.DataFrame) -> pd.DataFrame:
    df = df.copy()
    df["symptom_text_clean"] = df["symptom_text"].apply(sanitize_text)
    df["fix_text_clean"] = df["fix_text"].apply(sanitize_text)
    return df


# ---------------------------------------------------------------------------
# 3. VECTORIZE (turn text into numbers so a computer can compare notes)
# ---------------------------------------------------------------------------
def vectorize_text(texts):
    """
    TF-IDF = Term Frequency - Inverse Document Frequency.
    In plain English: it scores each word in a note by how important that
    word seems to be -- common words like "the"/"unit" get a low score,
    while distinctive words like "overheating" or "cracked" get a high
    score. Each note becomes a vector (a long list of numbers), and notes
    with similar important words end up as similar vectors.
    """
    n_docs = len(texts)

    # Custom stopwords: generic service-note boilerplate that shows up in
    # almost every note regardless of failure type ("customer", "reports",
    # "device", "unit"...). Left in, these words would dominate every
    # cluster centroid and drown out the words that actually distinguish
    # one failure type from another.
    #
    # Also includes generic vehicle nouns ("car", "vehicle", "model") and
    # filler adverbs/verbs ("properly", "keeps", "does") that showed up
    # in cluster labels like "Battery / Car / Properly" -- these say
    # nothing about the failure itself, they're just noise that happened
    # to have a high TF-IDF weight because the underlying dataset is
    # about cars. Same fix as the words already in this list, just
    # extended to cover what this dataset's notes actually look like.
    custom_stopwords = [
        "customer", "reports", "report", "reported", "device", "unit",
        "units", "issue", "problem", "time", "normal",
        "car", "cars", "vehicle", "vehicles", "model",
        "properly", "keeps", "keep", "does", "did", "doing",
        "always", "never", "completely", "really", "just", "quite",
    ]
    stopwords = list(TfidfVectorizer(stop_words="english").get_stop_words()) + custom_stopwords

    # min_df/max_df as fixed thresholds only make sense with a reasonably
    # large batch of notes (e.g. a big CSV/Excel export). A single
    # uploaded PDF or photo might only contain a handful of notes, where
    # "must appear in at least 2 documents" or "at most 40% of documents"
    # can conflict with each other and crash. So we relax these
    # constraints automatically for small uploads.
    if n_docs < 10:
        min_df, max_df = 1, 1.0
    else:
        min_df, max_df = 2, 0.4

    vectorizer = TfidfVectorizer(
        stop_words=stopwords,
        ngram_range=(1, 2),  # capture two-word phrases too, e.g. "no power"
        min_df=min_df,
        max_df=max_df,
        max_features=500,
    )
    try:
        matrix = vectorizer.fit_transform(texts)
    except ValueError:
        # A tiny or heavily redacted dataset can contain only stopwords.
        # Keep the analysis alive with a neutral token; the resulting single
        # cluster is more honest than crashing or inventing a theme.
        fallback_texts = [
            text if str(text).strip() else "failure"
            for text in texts
        ]
        fallback = TfidfVectorizer(ngram_range=(1, 2), min_df=1)
        matrix = fallback.fit_transform(fallback_texts)
        vectorizer = fallback
    return matrix, vectorizer


# ---------------------------------------------------------------------------
# 4. CLUSTER (group similar notes together)
# ---------------------------------------------------------------------------
def choose_best_k(matrix, k_range=range(4, 9)):
    """Choose k with a faster, still data-driven silhouette search.

    The old implementation could perform up to 50 KMeans initializations
    *per uploaded file*. Multi-file analysis now combines files first, so this
    search happens only once. We also use fewer KMeans restarts and, for very
    large datasets, score a representative sample rather than every row.
    """
    n_samples = matrix.shape[0]
    valid_range = [k for k in k_range if 2 <= k <= n_samples - 1]

    if not valid_range:
        if n_samples >= 3:
            return 2
        return 1

    # Keep the same 4-8 theme range, but don't spend forever scoring huge
    # uploads. Silhouette is only used to select k; the final KMeans still
    # runs on the complete dataset.
    sample_size = min(1000, n_samples)
    if n_samples > sample_size:
        from sklearn.model_selection import train_test_split
        sample_idx, _ = train_test_split(
            range(n_samples), train_size=sample_size, random_state=42
        )
        score_matrix = matrix[sample_idx]
    else:
        score_matrix = matrix

    best_k, best_score = valid_range[0], -1.0
    for k in valid_range:
        km = KMeans(n_clusters=k, random_state=42, n_init=3)
        labels = km.fit_predict(score_matrix)
        # A silhouette score needs at least two distinct labels.
        if len(set(labels)) < 2:
            continue
        score = silhouette_score(score_matrix, labels)
        if score > best_score:
            best_k, best_score = k, score
    return best_k


class _SingleClusterModel:
    """
    A tiny stand-in for a fitted KMeans model, used when there are too few
    notes to meaningfully cluster (e.g. 1-2 notes from a single uploaded
    PDF or photo). It just puts everything in one cluster, but still
    exposes `cluster_centers_` so label_clusters() can work on it exactly
    the same way it works on a real KMeans result -- no special-casing
    needed downstream.
    """
    def __init__(self, centroid):
        self.cluster_centers_ = [centroid]


def cluster_notes(matrix, n_clusters=None):
    n_samples = matrix.shape[0]

    if n_clusters is None:
        n_clusters = choose_best_k(matrix)

    # Can't ask KMeans for more clusters than we have notes -- cap it.
    n_clusters = max(1, min(n_clusters, n_samples))

    if n_clusters == 1:
        # Too few notes to cluster meaningfully -- put everything in one
        # group, labeled using the average of all the notes' keywords.
        centroid = matrix.mean(axis=0)
        centroid = centroid.A1 if hasattr(centroid, "A1") else centroid  # sparse -> dense array
        labels = [0] * n_samples
        km = _SingleClusterModel(centroid)
        return labels, km, 1

    km = KMeans(n_clusters=n_clusters, random_state=42, n_init=5)
    labels = km.fit_predict(matrix)
    return labels, km, n_clusters


# ---------------------------------------------------------------------------
# 5. LABEL CLUSTERS (turn "Cluster 3" into "Overheating / Thermal Shutdown")
# ---------------------------------------------------------------------------
def _clean_theme_label(top_terms, max_terms=3):
    """
    Turns a list of TF-IDF top terms (already sorted by weight, highest
    first) into a shorter, more readable label -- WITHOUT inventing any
    diagnosis that isn't in the terms themselves.

    Three purely mechanical cleanups, using only what TF-IDF already
    gave us (no synonym dictionary, no LLM):
      1. Prefer real two-word phrases over lone single words. vectorize_
         text() already captures these (ngram_range=(1, 2)) -- so if the
         notes actually contain a phrase like "battery start" or "start
         morning", that phrase is sitting right there in the same
         feature space, it just doesn't always out-rank single words on
         raw weight alone. Surfacing it first turns something like
         "Battery / Start / Morning" into "Battery Start / Morning" --
         still 100% derived from the real text, just phrased the way it
         actually appears in the notes instead of split into loose
         words.
      2. Skip a term if every one of its words is already covered by a
         term we've already picked (e.g. skip standalone "battery" once
         "battery start" was already picked) -- this removes the
         redundant keyword-soup look like "Battery / Car / Properly /
         Keeps" repeating the same idea.
      3. Within each group (phrases, then single words), keep the
         original weight-sorted order, so the label still reflects
         genuine TF-IDF importance rather than an arbitrary re-ranking.

    Caps at `max_terms` phrases so labels stay short enough for the UI.
    Falls back to the original top terms if this cleanup would leave
    nothing (shouldn't normally happen, but keeps this crash-proof).
    """
    phrases = [t for t in top_terms if " " in t]
    single_words = [t for t in top_terms if " " not in t]
    ordered_candidates = phrases + single_words

    selected = []
    used_words = set()

    for term in ordered_candidates:
        words = set(term.split())
        if words & used_words:
            continue
        selected.append(term)
        used_words |= words
        if len(selected) >= max_terms:
            break

    if not selected:
        selected = top_terms[:max_terms]

    return " / ".join(t.title() for t in selected)


def label_clusters(km, vectorizer, top_n=12, max_label_terms=3):
    """
    For each cluster, look at its centroid (the "average" note in that
    cluster) and pull out the words/phrases with the highest TF-IDF weight.
    Those words become the cluster's auto-generated label.

    We look at `top_n` candidate terms (more than we'll actually display)
    so that _clean_theme_label() has enough options both to find real
    2-word phrases among the candidates and to drop redundant,
    overlapping words -- purely keyword-based, no LLM involved.

    NOTE: In the original brief, this step was suggested to use Azure
    OpenAI to write a nicer label. That's a great upgrade -- just take
    these top keywords and ask an LLM: "these words describe a product
    failure theme, give it a short name." We keep it keyword-based here
    so the MVP needs no API key and is instant.
    """
    terms = vectorizer.get_feature_names_out()
    labels = {}
    for cluster_id, centroid in enumerate(km.cluster_centers_):
        top_indices = centroid.argsort()[::-1][:top_n]
        top_terms = [terms[i] for i in top_indices]
        labels[cluster_id] = _clean_theme_label(top_terms, max_terms=max_label_terms)
    return labels


# ---------------------------------------------------------------------------
# 5b. USE OPTIONAL STRUCTURED TAGS TO CLEAN UP / VALIDATE LABELS
# ---------------------------------------------------------------------------
# From the brief's "Data Considerations": some rows may already carry a
# structured tag like "no_power" or "overheating" (a technician picked it
# from a dropdown, say), but most won't -- real service data is
# inconsistent, which is exactly why we cluster in the first place instead
# of just grouping by tag.
#
# What we do with the tags that DO exist:
#   1. For each auto-discovered cluster, look at whichever of its notes
#      happen to have a tag, and find the most common one.
#   2. If that tag is used by a solid majority of the tagged notes in the
#      cluster ("agreement"), we trust the clustering and swap the raw
#      keyword-soup label (e.g. "Overheats / Fan / Hot / Shuts") for a
#      clean, human name built from the tag (e.g. "Overheating").
#   3. If tags disagree a lot, or there aren't enough tagged notes to
#      tell, we leave the keyword-based label alone -- it's still a
#      reasonable guess, we just don't have enough structured evidence to
#      override it.
# This also gives us a simple QA signal (tag_agreement_pct) worth
# surfacing in the insights: low agreement on a theme is a hint that the
# cluster might be mixing two different real-world failure modes.
def apply_structured_tags(df: pd.DataFrame, cluster_labels: dict, min_tagged=3, agreement_threshold=0.6):
    df = df.copy()
    tag_quality = {}

    for cluster_id, keyword_label in cluster_labels.items():
        cluster_rows = df[df["cluster_id"] == cluster_id]
        tagged = cluster_rows[cluster_rows["tag"] != ""]

        if len(tagged) >= min_tagged:
            top_tag, top_count = tagged["tag"].value_counts().idxmax(), tagged["tag"].value_counts().max()
            agreement = top_count / len(tagged)
        else:
            top_tag, agreement = None, 0.0

        if top_tag and agreement >= agreement_threshold:
            clean_name = top_tag.replace("_", " ").title()
            cluster_labels[cluster_id] = clean_name
        # else: keep the original keyword-based label as-is

        tag_quality[cluster_id] = {
            "tagged_notes": int(len(tagged)),
            "total_notes": int(len(cluster_rows)),
            "tag_agreement_pct": round(agreement * 100, 1),
            "used_tag_for_label": bool(top_tag and agreement >= agreement_threshold),
        }

    return cluster_labels, tag_quality


# ---------------------------------------------------------------------------
# FULL PIPELINE (glues steps 1-5 together)
# ---------------------------------------------------------------------------
def run_pipeline(path_or_buffer, filename: str = None, n_clusters=None):
    df = load_data(path_or_buffer, filename=filename)
    df = sanitize_dataframe(df)
    df = df[df["symptom_text_clean"].str.strip().ne("")].copy()
    if df.empty:
        raise ValueError("No usable complaint text was found in the uploaded file.")

    matrix, vectorizer = vectorize_text(df["symptom_text_clean"])
    cluster_ids, km, k_used = cluster_notes(matrix, n_clusters=n_clusters)
    cluster_labels = label_clusters(km, vectorizer)

    df["cluster_id"] = cluster_ids
    df["theme"] = df["cluster_id"].map(cluster_labels)

    # Use any optional structured tags to clean up / validate the labels
    # (see apply_structured_tags docstring above).
    cluster_labels, tag_quality = apply_structured_tags(df, cluster_labels)
    df["theme"] = df["cluster_id"].map(cluster_labels)
    df["tag_agreement_pct"] = df["cluster_id"].map(lambda c: tag_quality[c]["tag_agreement_pct"])

    return df, cluster_labels, k_used


# ---------------------------------------------------------------------------
# 6. TREND TABLE (counts per model per month per theme)
# ---------------------------------------------------------------------------
def run_pipeline_batch(file_items, n_clusters=None):
    """Run one complete analysis over multiple uploaded files.

    ``file_items`` is an iterable of ``(filename, bytes)`` pairs. Each file is
    parsed independently because formats differ, but all resulting rows are
    combined BEFORE sanitization, TF-IDF, clustering, and theme labeling.

    This preserves the existing features while making multiple uploads much
    faster and, importantly, gives all files a shared theme space.
    """
    from io import BytesIO

    frames = []
    detection_infos = []
    skipped_files = []

    for filename, file_bytes in file_items:
        try:
            df_i = load_data(BytesIO(file_bytes), filename=filename)
        except ValueError as exc:
            # One unrelated file should not prevent valid files in the same
            # multi-select upload from being analyzed.
            message = str(exc)
            if "not a supported input format" in message:
                skipped_files.append({"filename": filename, "reason": message})
                continue
            raise
        frames.append(df_i)
        detection_infos.append(
            df_i.attrs.get("detection_info", {"input_type": "structured"})
        )

    if not frames:
        if skipped_files:
            raise ValueError(
                "None of the uploaded files could be analyzed. "
                + " ".join(item["reason"] for item in skipped_files)
            )
        raise ValueError("No files were uploaded.")

    # Combine BEFORE the expensive ML steps.
    df = pd.concat(frames, ignore_index=True)
    df = sanitize_dataframe(df)
    df = df[df["symptom_text_clean"].str.strip().ne("")].copy()
    if df.empty:
        raise ValueError("No usable complaint text was found in the uploaded files.")

    matrix, vectorizer = vectorize_text(df["symptom_text_clean"])
    cluster_ids, km, k_used = cluster_notes(matrix, n_clusters=n_clusters)
    cluster_labels = label_clusters(km, vectorizer)

    df["cluster_id"] = cluster_ids
    df["theme"] = df["cluster_id"].map(cluster_labels)

    cluster_labels, tag_quality = apply_structured_tags(df, cluster_labels)
    df["theme"] = df["cluster_id"].map(cluster_labels)
    df["tag_agreement_pct"] = df["cluster_id"].map(
        lambda c: tag_quality[c]["tag_agreement_pct"]
    )

    trend = build_trend_table(df)

    raw_info = next(
        (info for info in detection_infos if info.get("input_type") == "raw_reviews"),
        None,
    )
    detection_info = raw_info or {
        "input_type": "structured",
        "files_processed": len(frames),
    }
    if raw_info is not None:
        detection_info = dict(raw_info)
        detection_info["files_processed"] = len(frames)
    if skipped_files:
        detection_info = dict(detection_info)
        detection_info["skipped_files"] = skipped_files

    df.attrs["detection_info"] = detection_info
    return df, cluster_labels, k_used, trend, detection_info


def build_trend_table(df: pd.DataFrame) -> pd.DataFrame:
    df = df.copy()
    # Missing dates are valid for ordinary text-only exports.  Do not turn
    # them into fabricated "today" values or a misleading temporal trend.
    df["date"] = pd.to_datetime(df["date"], errors="coerce")
    df = df[df["date"].notna()]
    if df.empty:
        return pd.DataFrame(columns=["month", "product_model", "theme", "count"])
    df["month"] = df["date"].dt.to_period("M").astype(str)
    trend = (
        df.groupby(["month", "product_model", "theme"])
        .size()
        .reset_index(name="count")
    )
    return trend


# ---------------------------------------------------------------------------
# 7. FAILURE PRIORITY / COMPONENT / INVESTIGATION HELPERS
# ---------------------------------------------------------------------------
_COMPONENT_HINTS = {
    "Battery": ("battery", "low charge", "won't start", "wont start", "no start"),
    "Starter": ("starter", "clicking", "crank", "cranking"),
    "Alternator": ("alternator", "charging system", "not charging"),
    "Brakes": ("brake", "braking", "rotor", "pad"),
    "Engine": ("engine", "motor", "stall", "misfire", "rpm"),
    "AC Compressor": ("compressor", "air conditioning", "a/c", "ac ", "cooling"),
    "Cooling System": ("overheat", "overheating", "radiator", "coolant", "fan"),
    "Fuel Pump": ("fuel pump", "fuel pressure", "fuel delivery"),
    "Electrical": ("wiring", "fuse", "short circuit", "electrical", "power loss"),
    "Display / Screen": ("screen", "display", "touchscreen"),
}

_SERIOUS_FAILURE_TERMS = (
    "fire", "smoke", "crash", "accident", "brake failure", "brakes failed",
    "stall", "overheat", "overheating", "leak", "unsafe", "injury",
)


def _likely_component(rows: pd.DataFrame) -> str:
    if "component" in rows.columns:
        known = rows["component"].fillna("").astype(str).str.strip()
        known = known[~known.str.lower().isin({"", "unknown", "nan", "none"})]
        if len(known):
            return str(known.value_counts().idxmax())
    text = " ".join(rows.get("symptom_text_clean", rows["symptom_text"])
                    .fillna("").astype(str)).lower()
    scores = {
        component: sum(text.count(term) for term in terms)
        for component, terms in _COMPONENT_HINTS.items()
    }
    best = max(scores, key=scores.get) if scores else None
    return best if best and scores[best] else "Not identified"


def build_failure_analysis(df: pd.DataFrame, trend: pd.DataFrame = None) -> dict:
    """Return transparent, lightweight cluster-level decision support.

    This is deliberately a heuristic, not a model confidence score.  The
    returned values are safe to show beside the existing TF-IDF/KMeans output
    and also work when optional columns are absent.
    """
    if df is None or df.empty:
        return {}
    total = len(df)
    trend = trend if trend is not None else build_trend_table(df)
    valid_dates = pd.to_datetime(df.get("date"), errors="coerce").notna().sum()
    result = {}
    for theme, rows in df.groupby("theme", sort=False):
        count = len(rows)
        pct = 100 * count / total if total else 0
        text = " ".join(rows.get("symptom_text_clean", rows["symptom_text"])
                        .fillna("").astype(str)).lower()
        serious_hits = sum(text.count(term) for term in _SERIOUS_FAILURE_TERMS)
        recent_change = None
        if valid_dates and trend is not None and not trend.empty:
            counts = trend[trend["theme"] == theme].groupby("month")["count"].sum()
            if len(counts) >= 2:
                midpoint = max(1, len(counts) // 2)
                old, new = counts.iloc[:midpoint].sum(), counts.iloc[midpoint:].sum()
                if old:
                    recent_change = round(100 * (new - old) / old, 1)
        if serious_hits or pct >= 35:
            priority = "Critical"
        elif pct >= 20 or (recent_change is not None and recent_change >= 50):
            priority = "High"
        elif pct >= 8 or (recent_change is not None and recent_change > 0):
            priority = "Medium"
        else:
            priority = "Low"

        symptoms = _select_representative_examples(
            rows["symptom_text_clean"], n=3
        ) if "symptom_text_clean" in rows else []
        component = _likely_component(rows)
        if recent_change is None:
            trend_label = "Unavailable (no usable date trend)"
        elif recent_change > 10:
            trend_label = f"Increasing ({recent_change:+g}% recently)"
        elif recent_change < -10:
            trend_label = f"Decreasing ({recent_change:+g}% recently)"
        else:
            trend_label = "Stable"

        cause = f"{component} may be contributing to the reported symptoms." \
            if component != "Not identified" else \
            "The available complaint text does not identify a clear component."
        action = f"Inspect {component.lower()} and review related repair records." \
            if component != "Not identified" else \
            "Review representative complaints and collect component/repair details."
        five_whys = [
            f"Complaints report {str(theme).lower()} symptoms.",
            f"The reported symptoms may involve {component.lower()}.",
            "The component may be worn, disconnected, or out of specification.",
            "A related upstream condition may be contributing to the failure.",
            "Inspect the component and confirm the cause against repair evidence.",
        ]
        result[str(theme)] = {
            "count": int(count), "pct_of_total": round(pct, 1),
            "priority": priority, "component": component,
            "trend": trend_label, "trend_change_pct": recent_change,
            "symptoms": symptoms, "possible_root_cause": cause,
            "recommended_action": action,
            "five_whys": five_whys,
            "temporal_analysis_available": bool(valid_dates and not trend.empty),
        }
    return result


# ---------------------------------------------------------------------------
# 7. INSIGHTS SUMMARY (plain-English top drivers + examples)
# ---------------------------------------------------------------------------
def _select_representative_examples(texts, n=3):
    """
    Picks the `n` notes that are most representative of a failure theme,
    using the SAME TF-IDF representation the pipeline already builds
    everything else from (see vectorize_text()) -- no new algorithm, no
    embeddings, no LLM.

    "Representative" here means: closest (highest cosine similarity) to
    the centroid of this theme's own notes in TF-IDF space -- exactly the
    same idea KMeans uses internally (a centroid is just the mean of the
    vectors assigned to it). This replaces the previous behavior, which
    just took whichever unique notes happened to appear first, and could
    surface an outlier note that doesn't actually support the theme.

    Falls back to the first `n` notes (the old behavior) if there are too
    few notes to vectorize meaningfully, or if anything about the TF-IDF
    step fails -- this must never crash the Insights page.
    """
    unique_texts = pd.Series(texts).drop_duplicates().tolist()

    if len(unique_texts) <= n:
        return unique_texts

    try:
        matrix, _ = vectorize_text(unique_texts)
        centroid = np.asarray(matrix.mean(axis=0))
        similarities = cosine_similarity(matrix, centroid).ravel()
        top_indices = similarities.argsort()[::-1][:n]
        return [unique_texts[i] for i in top_indices]
    except Exception:
        return unique_texts[:n]


def build_insights(df: pd.DataFrame, top_n=5, examples_per_theme=3):
    theme_counts = df["theme"].value_counts().head(top_n)

    insights = []
    for theme, count in theme_counts.items():
        pct = round(100 * count / len(df), 1)
        theme_rows = df[df["theme"] == theme]
        examples = _select_representative_examples(
            theme_rows["symptom_text_clean"],
            n=examples_per_theme,
        )
        top_model = theme_rows["product_model"].value_counts().idxmax()

        # Key terms for the "Why this pattern?" section -- reuses the
        # theme's own auto-generated label (already TF-IDF-derived, see
        # label_clusters()/_clean_theme_label()) rather than recomputing
        # anything, so this never claims more than the existing model
        # actually knows.
        key_terms = []
        for phrase in str(theme).split(" / "):
            for word in phrase.split():
                word = word.strip().lower()
                if word and word not in key_terms:
                    key_terms.append(word)

        # Serial range(s) most represented in this theme -- helps spot
        # whether a failure is concentrated in one manufacturing batch.
        serial_counts = theme_rows["serial_range"].value_counts()
        top_serial = serial_counts.idxmax() if len(serial_counts) else "Unknown"

        tag_agreement = (
            round(theme_rows["tag_agreement_pct"].iloc[0], 1)
            if "tag_agreement_pct" in theme_rows.columns and len(theme_rows)
            else 0.0
        )

        insights.append({
            "theme": theme,
            "count": int(count),
            "pct_of_total": pct,
            "most_affected_model": top_model,
            "most_common_serial_range": top_serial,
            "tag_agreement_pct": tag_agreement,
            "key_terms": key_terms,
            "examples": examples,
        })
    return insights


# ---------------------------------------------------------------------------
# 8. WORD REPORT (one downloadable .docx with charts + summary)
# ---------------------------------------------------------------------------
# WHY MATPLOTLIB HERE INSTEAD OF THE DASHBOARD'S PLOTLY CHARTS?
# The dashboard's charts (in app.py) are interactive Plotly charts, which
# is great for browsing on screen. But turning a Plotly chart into a
# static image for a Word doc needs an extra tool (Kaleido) that in turn
# needs Google Chrome installed on the computer -- one more install step
# that can trip people up. Matplotlib draws directly to an image with no
# extra dependencies, so we use it just for this report, while the
# on-screen dashboard keeps its nicer interactive Plotly charts.

def _chart_theme_bar(df: pd.DataFrame):
    """Horizontal bar chart of note count per theme -> returns a PNG buffer."""
    import matplotlib
    matplotlib.use("Agg")  # no GUI needed, just render to an image
    import matplotlib.pyplot as plt

    counts = df["theme"].value_counts().sort_values()
    fig, ax = plt.subplots(figsize=(7, 4))
    ax.barh(counts.index, counts.values, color="#2E74B5")
    ax.set_xlabel("Number of Notes")
    ax.set_title("Failure Themes (Overall)")
    fig.tight_layout()

    buf = io.BytesIO()
    fig.savefig(buf, format="png", dpi=150)
    plt.close(fig)
    buf.seek(0)
    return buf


def _chart_trend_line(trend: pd.DataFrame):
    """Line chart of note count per month, one line per theme -> PNG buffer."""
    import matplotlib
    matplotlib.use("Agg")
    import matplotlib.pyplot as plt

    pivot = trend.groupby(["month", "theme"])["count"].sum().unstack(fill_value=0)
    pivot = pivot.sort_index()

    fig, ax = plt.subplots(figsize=(7, 4))
    for theme in pivot.columns:
        ax.plot(pivot.index, pivot[theme], marker="o", label=theme, linewidth=1.5)
    ax.set_xlabel("Month")
    ax.set_ylabel("Number of Notes")
    ax.set_title("Trend Over Time by Theme")
    ax.tick_params(axis="x", rotation=45)
    ax.legend(fontsize=7, loc="upper left", bbox_to_anchor=(1.02, 1))
    fig.tight_layout()

    buf = io.BytesIO()
    fig.savefig(buf, format="png", dpi=150)
    plt.close(fig)
    buf.seek(0)
    return buf


def _chart_model_stack(df: pd.DataFrame):
    """Stacked bar chart of failure counts per product model -> PNG buffer."""
    import matplotlib
    matplotlib.use("Agg")
    import matplotlib.pyplot as plt

    pivot = df.groupby(["product_model", "theme"]).size().unstack(fill_value=0)

    fig, ax = plt.subplots(figsize=(7, 4))
    bottom = None
    for theme in pivot.columns:
        ax.bar(pivot.index, pivot[theme], bottom=bottom, label=theme)
        bottom = pivot[theme] if bottom is None else bottom + pivot[theme]
    ax.set_ylabel("Number of Notes")
    ax.set_title("Failure Counts by Product Model")
    ax.legend(fontsize=7, loc="upper left", bbox_to_anchor=(1.02, 1))
    fig.tight_layout()

    buf = io.BytesIO()
    fig.savefig(buf, format="png", dpi=150)
    plt.close(fig)
    buf.seek(0)
    return buf


def build_word_report(
    df: pd.DataFrame,
    trend: pd.DataFrame,
    insights: list,
    k_used: int,
    source_name: str = "uploaded data",
) -> bytes:
    """
    Builds a single Word (.docx) report covering everything the dashboard
    shows: a summary, all three charts (as images), and the actionable
    insights with example notes. Returns the file as raw bytes, ready to
    hand to Streamlit's download_button.
    """
    from docx import Document
    from docx.shared import Inches, Pt, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH

    NAVY = RGBColor(0x1F, 0x38, 0x64)

    doc = Document()

    # --- Title ---
    title = doc.add_heading("Field Returns Failure Mode Mining", level=0)
    title.runs[0].font.color.rgb = NAVY

    sub = doc.add_paragraph()
    sub.add_run(f"Analysis Report  •  Source: {source_name}  •  "
                 f"Generated: {datetime.today().strftime('%d %b %Y')}").italic = True

    # --- Executive summary / KPIs ---
    doc.add_heading("Summary", level=1)
    top_theme = df["theme"].value_counts().idxmax() if len(df) else "N/A"
    kpi_table = doc.add_table(rows=1, cols=4)
    kpi_table.style = "Light Grid Accent 1"
    hdr = kpi_table.rows[0].cells
    kpi_labels = ["Total Notes", "Themes Discovered", "Top Failure Mode", "Product Models"]
    kpi_values = [str(len(df)), str(k_used), top_theme, str(df["product_model"].nunique())]
    for i, label in enumerate(kpi_labels):
        hdr[i].text = label
    row2 = kpi_table.add_row().cells
    for i, val in enumerate(kpi_values):
        row2[i].text = val

    doc.add_paragraph()

    # --- Charts ---
    doc.add_heading("Failure Themes (Overall)", level=1)
    doc.add_picture(_chart_theme_bar(df), width=Inches(6))

    doc.add_heading("Trend Over Time by Theme", level=1)
    doc.add_picture(_chart_trend_line(trend), width=Inches(6))

    doc.add_heading("Failure Counts by Product Model", level=1)
    doc.add_picture(_chart_model_stack(df), width=Inches(6))

    # --- Actionable insights ---
    doc.add_heading("Actionable Insights Summary", level=1)
    for ins in insights:
        h = doc.add_heading(
            f"{ins['theme']}  —  {ins['count']} notes "
            f"({ins['pct_of_total']}% of total)",
            level=2,
        )
        doc.add_paragraph(f"Most affected model: {ins['most_affected_model']}")

        serial_range = ins.get("most_common_serial_range")
        if serial_range and serial_range != "Unknown":
            doc.add_paragraph(f"Most common serial range: {serial_range}")

        if ins.get("tag_agreement_pct", 0) > 0:
            doc.add_paragraph(
                f"Structured-tag agreement: {ins['tag_agreement_pct']}% "
                "(how often the technician's own tag matched this auto-discovered theme)"
            )

        if ins.get("key_terms"):
            doc.add_paragraph(f"Key terms: {' · '.join(ins['key_terms'])}")

        doc.add_paragraph("Example notes:")
        for ex in ins["examples"]:
            doc.add_paragraph(ex, style="List Bullet")

    # --- Footer note ---
    doc.add_paragraph()
    note = doc.add_paragraph()
    note.add_run(
        "MVP scope only — production hardening (robust PII detection, "
        "human-in-the-loop labeling, live data ingestion) not included."
    ).italic = True

    buffer = io.BytesIO()
    doc.save(buffer)
    return buffer.getvalue()


if __name__ == "__main__":
    # Quick manual test: run the whole pipeline and print a summary.
    df, labels, k = run_pipeline("data/service_notes.csv")
    print(f"Chosen number of clusters (k): {k}")
    print("\nDiscovered themes:")
    for cid, label in labels.items():
        count = (df["cluster_id"] == cid).sum()
        print(f"  Cluster {cid}: {label}  ({count} notes)")

    print("\nSample sanitized notes (check PII removed):")
    pii_mask = df["symptom_text"].str.contains(r"@|\+91|Customer [A-Z]", regex=True, na=False)
    for _, sample in df[pii_mask].head(3).iterrows():
        print("  BEFORE:", sample["symptom_text"])
        print("  AFTER :", sample["symptom_text_clean"])
        print()